"""공통 Document 서비스 — 도메인 비종속 (설계 §4, Critical Design Rule #2).

소유 주체는 범용 `owner_id` + `domain` 으로만 표현한다(bot_id 미사용). 영수증 컴플라이언스
근거 문서는 `domain="expense_rule" AND is_compliance_source=true` 로 적재되어, 일반 규정
(domain="policy")과 검색 단계에서 완전 분리된다(Rule #1).

삭제는 Hard Delete 하지 않고 `embedding_status="DELETING"` 으로 전이만 한다(Rule #4):
조회/목록에서 즉시 제외되어 '좀비 벡터' 재노출을 막고, 실제 물리 정리(벡터→파일→행)는
백그라운드 워커가 멱등 수행한다(별도 페이즈).

⚠️ 스캐폴딩 범위: 비동기/동기 임베딩(파싱·청킹·벡터 적재)은 Phase B 에서 구현한다. 본
파일은 문서 행 생성·파일 저장·목록/단건/다운로드·Soft Delete 까지를 동작시키고, 실제
벡터화는 TODO 로 표시한다(상태는 PROCESSING 유지).
"""
from __future__ import annotations

import logging
import os

from langchain_text_splitters import RecursiveCharacterTextSplitter
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.ai.vector_store import get_policy_vector_store
from app.core.dependencies import TenantContext
from app.db.session import SessionLocal
from app.models.documents import Document
from app.schemas.documents import DocumentIngestTextRequest

logger = logging.getLogger(__name__)

# 임베딩 라이프사이클 상태값(설계 §2.2).
STATUS_PROCESSING = "PROCESSING"
STATUS_COMPLETED = "COMPLETED"
STATUS_FAILED = "FAILED"
STATUS_DELETING = "DELETING"

# 청킹 파라미터(설계 §4.2 — policy_service 와 동일한 보수적 설정).
_CHUNK_SIZE = 500
_CHUNK_OVERLAP = 50

# 업로드 파일 저장 루트(테넌트 격리 경로의 베이스). TODO: 추후 settings 로 이동 고려.
UPLOAD_BASE_DIR = os.environ.get("UPLOAD_DIR", "uploads")


def _extract_text(file_path: str, content_type: str | None, file_name: str | None) -> str:
    """파일에서 텍스트 추출. PDF=pypdf, DOCX=python-docx, 그 외=UTF-8 텍스트.

    파서 라이브러리(pypdf/python-docx)는 lazy import 한다 — 미설치 시 해당 형식만 실패하고
    텍스트(.txt/.md 등)는 외부 의존성 없이 동작한다(설계: "유틸 없으면 기본 처리").
    """
    name = (file_name or "").lower()
    ctype = (content_type or "").lower()
    is_pdf = name.endswith(".pdf") or "pdf" in ctype
    is_docx = name.endswith(".docx") or "wordprocessingml" in ctype

    if is_pdf:
        from pypdf import PdfReader  # lazy import

        reader = PdfReader(file_path)
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if is_docx:
        import docx  # python-docx, lazy import

        document = docx.Document(file_path)
        return "\n".join(p.text for p in document.paragraphs)
    # 기본: 텍스트로 읽기(인코딩 관용 — 깨진 바이트는 대체)
    with open(file_path, "r", encoding="utf-8", errors="replace") as fh:
        return fh.read()


def _split_text(text: str) -> list[str]:
    """RecursiveCharacterTextSplitter(500/50)로 청킹."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=_CHUNK_SIZE, chunk_overlap=_CHUNK_OVERLAP
    )
    return splitter.split_text(text)


def _build_chunk_metadatas(doc: Document, n_chunks: int) -> list[dict]:
    """⚠️ Qdrant payload 메타데이터 — 격리/필터 필드를 **반드시** 포함(설계 §3, Rule #1).

    company_id/workplace_id(테넌트 격리), domain/is_compliance_source(컴플라이언스 이중게이트),
    owner_id(소유), doc_id(삭제 단위), source/title/chunk_index(표시·출처).
    """
    return [
        {
            "company_id": doc.company_id,
            "workplace_id": doc.workplace_id,
            "domain": doc.domain,
            "is_compliance_source": doc.is_compliance_source,
            "owner_id": doc.owner_id,
            "doc_id": doc.id,
            "source": doc.source_name,
            "title": doc.title,
            "chunk_index": i,
        }
        for i in range(n_chunks)
    ]


def _embed_and_store(doc: Document, text: str, vector_store=None) -> int:
    """청킹 → 메타데이터 부여 임베딩 → tenant_documents 컬렉션 적재. 반환: 적재된 청크 수.

    빈 텍스트(청크 0개)면 적재 없이 0 을 반환(성공 처리). vector_store 미지정 시 공용
    싱글톤(get_policy_vector_store)을 쓴다 — 테스트는 in-memory 스토어를 주입한다.
    """
    chunks = _split_text(text)
    if not chunks:
        return 0
    vs = vector_store or get_policy_vector_store()
    vs.add_texts(texts=chunks, metadatas=_build_chunk_metadatas(doc, len(chunks)))
    return len(chunks)


def process_document_embedding(doc_id: str, *, vector_store=None) -> None:
    """[Background] 업로드 문서 비동기 임베딩 파이프라인 (FastAPI BackgroundTasks).

    요청 세션이 닫힌 뒤 실행되므로 내부에서 **새 SessionLocal()** 을 연다.
    흐름: 파일 파싱 → 청킹 → 메타데이터 부여 임베딩 → tenant_documents 적재 →
    status=COMPLETED(+chunk_count). 예외 시 status=FAILED(+error_message).
    삭제 중(DELETING)이면 즉시 스킵(삭제 우선 — 좀비 벡터 방지).
    """
    db = SessionLocal()
    try:
        doc = db.get(Document, doc_id)
        if doc is None:
            logger.warning("임베딩 스킵 — 문서 없음: %s", doc_id)
            return
        if doc.embedding_status == STATUS_DELETING:
            logger.info("임베딩 스킵 — 삭제 중(DELETING): %s", doc_id)
            return
        try:
            doc.embedding_status = STATUS_PROCESSING  # 재확인(이미 PROCESSING)
            text = (
                _extract_text(doc.file_path, doc.content_type, doc.file_name)
                if doc.file_path
                else ""
            )
            count = _embed_and_store(doc, text, vector_store=vector_store)
            doc.embedding_status = STATUS_COMPLETED
            doc.chunk_count = count
            doc.error_message = None
            logger.info("임베딩 완료: doc_id=%s chunks=%d", doc_id, count)
        except Exception as exc:  # noqa: BLE001 -- 실패는 FAILED 로 기록하고 워커는 죽지 않음
            doc.embedding_status = STATUS_FAILED
            doc.error_message = str(exc)[:1024]
            logger.exception("임베딩 실패 -> FAILED: doc_id=%s", doc_id)
        db.commit()
    finally:
        db.close()


class DocumentNotFoundError(Exception):
    """문서가 없거나 현재 테넌트 소유가 아니거나 DELETING 상태일 때(존재 비노출 위해 404)."""

    def __init__(self, doc_id: str) -> None:
        self.doc_id = doc_id
        super().__init__(f"document not found: {doc_id}")


class DocumentService:
    """공통 문서 메타/파일/Soft Delete 도메인 서비스.

    vector_store 는 Phase B 에서 임베딩에 사용하기 위해 lazy 주입 가능하게 둔다(현재 미사용).
    """

    def __init__(self, db: Session, vector_store=None) -> None:
        self.db = db
        self._vector_store = vector_store  # Phase B 임베딩용(미사용)

    # ------------------------------------------------------------------ #
    # Helpers
    # ------------------------------------------------------------------ #
    def _get_owned(self, doc_id: str, tenant: TenantContext) -> Document:
        """현재 테넌트 소유 + 비-DELETING 문서 조회. 없으면 DocumentNotFoundError(404)."""
        stmt = select(Document).where(
            Document.id == doc_id,
            Document.company_id == tenant.company_id,
            Document.workplace_id == tenant.workplace_id,
            Document.embedding_status != STATUS_DELETING,
        )
        doc = self.db.execute(stmt).scalar_one_or_none()
        if doc is None:
            raise DocumentNotFoundError(doc_id)
        return doc

    # ------------------------------------------------------------------ #
    # List / Get
    # ------------------------------------------------------------------ #
    def list_documents(
        self,
        tenant: TenantContext,
        domain: str | None = None,
        owner_id: str | None = None,
    ) -> list[Document]:
        """테넌트 + (선택)domain/owner_id 필터 문서 목록(DELETING 제외, 최신순)."""
        filters = [
            Document.company_id == tenant.company_id,
            Document.workplace_id == tenant.workplace_id,
            Document.embedding_status != STATUS_DELETING,
        ]
        if domain is not None:
            filters.append(Document.domain == domain)
        if owner_id is not None:
            filters.append(Document.owner_id == owner_id)
        stmt = select(Document).where(*filters).order_by(Document.created_at.desc())
        return list(self.db.execute(stmt).scalars().all())

    def get_document(self, doc_id: str, tenant: TenantContext) -> Document:
        return self._get_owned(doc_id, tenant)

    def get_for_download(self, doc_id: str, tenant: TenantContext) -> Document:
        """다운로드 대상 문서 조회. 파일이 없으면 DocumentNotFoundError(404)."""
        doc = self._get_owned(doc_id, tenant)
        if not doc.file_path or not os.path.isfile(doc.file_path):
            raise DocumentNotFoundError(doc_id)
        return doc

    # ------------------------------------------------------------------ #
    # Create (upload / ingest-text)
    # ------------------------------------------------------------------ #
    def create_from_upload(
        self,
        tenant: TenantContext,
        *,
        file,  # fastapi.UploadFile (동기 스택이므로 file.file 로 읽는다)
        title: str,
        owner_id: str | None,
        domain: str,
        is_compliance_source: bool,
    ) -> Document:
        """파일 업로드 → 테넌트 격리 경로 저장 + 문서 행 생성(status=PROCESSING).

        실제 임베딩은 라우터가 BackgroundTasks 로 process_document_embedding 를 예약한다
        (Phase B 에서 구현). 본 메서드는 파일 저장과 행 생성까지만 동작한다.
        """
        doc = Document(
            company_id=tenant.company_id,
            workplace_id=tenant.workplace_id,
            domain=domain,
            owner_id=owner_id,
            is_compliance_source=is_compliance_source,
            title=title,
            file_name=file.filename,
            content_type=file.content_type,
            source_name=title,
            embedding_status=STATUS_PROCESSING,
        )
        self.db.add(doc)
        self.db.flush()  # doc.id 확보(저장 경로에 사용)

        dest_dir = os.path.join(
            UPLOAD_BASE_DIR, tenant.company_id, tenant.workplace_id, doc.id
        )
        os.makedirs(dest_dir, exist_ok=True)
        dest_path = os.path.join(dest_dir, file.filename or doc.id)
        content = file.file.read()  # 동기 읽기(SpooledTemporaryFile)
        with open(dest_path, "wb") as fh:
            fh.write(content)

        doc.file_path = dest_path
        doc.byte_size = len(content)
        self.db.commit()
        self.db.refresh(doc)
        logger.info(
            "문서 업로드: tenant=%s/%s id=%s domain=%s owner=%s compliance=%s file=%r",
            tenant.company_id, tenant.workplace_id, doc.id, domain, owner_id,
            is_compliance_source, file.filename,
        )
        return doc

    def ingest_text(
        self, tenant: TenantContext, payload: DocumentIngestTextRequest
    ) -> Document:
        """텍스트 즉시 적재(동기) — 행 생성 후 그 자리에서 청킹/임베딩 → COMPLETED(설계 §4.2-4).

        예외 시 FAILED(+error_message). vector_store 는 주입값(self._vector_store) 우선,
        없으면 공용 싱글톤(get_policy_vector_store)을 사용한다.
        """
        doc = Document(
            company_id=tenant.company_id,
            workplace_id=tenant.workplace_id,
            domain=payload.domain,
            owner_id=payload.owner_id,
            is_compliance_source=payload.is_compliance_source,
            title=payload.source_name,
            file_name=None,
            source_name=payload.source_name,
            embedding_status=STATUS_PROCESSING,
        )
        self.db.add(doc)
        self.db.flush()  # doc.id 확보(metadata.doc_id 에 사용)
        try:
            count = _embed_and_store(doc, payload.text, vector_store=self._vector_store)
            doc.embedding_status = STATUS_COMPLETED
            doc.chunk_count = count
            doc.error_message = None
        except Exception as exc:  # noqa: BLE001 -- 실패는 FAILED 로 기록
            doc.embedding_status = STATUS_FAILED
            doc.error_message = str(exc)[:1024]
            logger.exception("텍스트 적재 임베딩 실패 -> FAILED: doc_id=%s", doc.id)
        self.db.commit()
        self.db.refresh(doc)
        logger.info(
            "문서 텍스트 적재: tenant=%s/%s id=%s domain=%s owner=%s compliance=%s status=%s",
            tenant.company_id, tenant.workplace_id, doc.id,
            payload.domain, payload.owner_id, payload.is_compliance_source, doc.embedding_status,
        )
        return doc

    # ------------------------------------------------------------------ #
    # Soft Delete (§4.3)
    # ------------------------------------------------------------------ #
    def soft_delete(self, doc_id: str, tenant: TenantContext) -> Document:
        """Soft Delete — embedding_status="DELETING" 으로 전이만 한다(Rule #4).

        즉시 조회/검색/목록에서 제외된다. 실제 물리 정리(벡터 doc_id 필터 멱등 delete →
        디스크 파일 → DB 행)는 백그라운드 워커가 수행한다(별도 페이즈).
        """
        doc = self._get_owned(doc_id, tenant)
        doc.embedding_status = STATUS_DELETING
        self.db.commit()
        self.db.refresh(doc)
        logger.info(
            "문서 Soft Delete(status=DELETING): tenant=%s/%s id=%s",
            tenant.company_id, tenant.workplace_id, doc.id,
        )
        # TODO(worker, 별도 페이즈): 벡터(metadata.doc_id 필터, 멱등) → 파일 → DB 행 순 제거.
        return doc
